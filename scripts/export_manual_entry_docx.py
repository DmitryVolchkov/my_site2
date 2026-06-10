"""Конвертация docs/manual-entry/README.md в Word (.docx) с картинками."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MANUAL_DIR = ROOT / "docs" / "manual-entry"
SOURCE = MANUAL_DIR / "README.md"
OUTPUT = MANUAL_DIR / "instrukciya-ruchnogo-vvoda.docx"

IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^- (.*)$")
CHECKBOX_RE = re.compile(r"^- \[([ xX])\] (.*)$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_rich_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            cursor = 0
            for match in LINK_RE.finditer(part):
                if match.start() > cursor:
                    paragraph.add_run(part[cursor : match.start()])
                link_run = paragraph.add_run(match.group(1))
                link_run.underline = True
                cursor = match.end()
            if cursor < len(part):
                paragraph.add_run(part[cursor:])


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "---") for cell in cells)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = cell
    document.add_paragraph()


def add_image(document: Document, base_dir: Path, alt: str, rel_path: str) -> None:
    image_path = (base_dir / rel_path).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Изображение не найдено: {rel_path}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    if alt:
        caption = document.add_paragraph(alt)
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        caption.runs[0].italic = True


def convert(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_heading(lines[0].lstrip("# ").strip(), level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    idx = 1
    while idx < len(lines):
        line = lines[idx]
        idx += 1

        if not line.strip():
            continue
        if line.strip() == "---":
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = len(heading.group(1))
            document.add_heading(heading.group(2).strip(), level=level)
            continue

        image = IMAGE_RE.match(line)
        if image:
            add_image(document, source.parent, image.group(1), image.group(2))
            continue

        if line.strip().startswith("```"):
            code_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            if idx < len(lines):
                idx += 1
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            continue

        if line.strip().startswith("|"):
            table_rows: list[list[str]] = []
            while idx - 1 < len(lines) and lines[idx - 1].strip().startswith("|"):
                row_line = lines[idx - 1]
                if not is_table_separator(row_line):
                    table_rows.append(parse_table_row(row_line))
                idx += 1
            add_table(document, table_rows)
            continue

        if line.startswith(">"):
            paragraph = document.add_paragraph(style="Intense Quote")
            add_rich_text(paragraph, line.lstrip("> ").strip())
            continue

        checkbox = CHECKBOX_RE.match(line)
        if checkbox:
            mark = "☑" if checkbox.group(1).lower() == "x" else "☐"
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, f"{mark} {checkbox.group(2)}")
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, ordered.group(2))
            continue

        unordered = UNORDERED_RE.match(line)
        if unordered:
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, unordered.group(1))
            continue

        paragraph = document.add_paragraph()
        add_rich_text(paragraph, line.strip())

    document.save(output)
    print(f"OK: {output}")


def main() -> int:
    if not SOURCE.exists():
        print(f"Файл не найден: {SOURCE}", file=sys.stderr)
        return 1
    convert(SOURCE, OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
